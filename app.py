Professional CV Builder Web Application
Built with Streamlit - Modern, Clean, User-Friendly Interface
Supports PDF and Word export with live preview and Turkish character support
"""

import streamlit as st
import io
import os
from datetime import datetime, date
from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import inch
from reportlab.lib import colors
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle
from reportlab.pdfbasment
from docx.shared import Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

# Page configuration
st.set_page_config(
    page_title="Professional CV Builder",
    page_icon="📄",
    layout="wide",
    initial_sidebar_state="expanded"
)

# Custom CSS for modern styling
st.markdown("""
<style>
    .main-header {
        font-size: 3rem;
        font-weight: 700;
        color: #1f77b4;
        text-align: center;
        margin-bottom: 2rem;
        text-shadow: 2px 2px 4px rgba(0,0,0,0.1);
    }
    
    .section-header {
        font-size: 1.5rem;
        font-weight: 600;
        color: #2c3e50;
        margin: 1.5rem 0 1rem 0;
        padding-bottom: 0.5rem;
        border-bottom: 2px solid #3498db;
    }
    
    .info-box {
        background: linear-gradient(135deg, #667eea 0%, #764ba2 100%);
        padding: 1rem;
        border-radius: 10px;
        color: white;
        margin: 1rem 0;
    }
    
    .success-box {
        background: linear-gradient(135deg, #4facfe 0%, #00f2fe 100%);
        padding: 1rem;
        border-radius: 10px;
        color: white;
        margin: 1rem 0;
    }
    
    .warning-box {
        background: linear-gradient(135deg, #fa709a 0%, #fee140 100%);
        padding: 1rem;
        border-radius: 10px;
        color: white;
        margin: 1rem 0;
    }
    
    .preview-container {
        background: #f8f9fa;
        padding: 2rem;
        border-radius: 10px;
        border: 1px solid #dee2e6;
        margin: 1rem 0;
    }
    
    .download-button {
        background: linear-gradient(135deg, #667eea 0%, #764ba2 100%);
        color: white;
        border: none;
        padding: 0.75rem 2rem;
        border-radius: 25px;
        font-weight: 600;
        margin: 0.5rem;
    }
    
    .stButton > button {
        background: linear-gradient(135deg, #667eea 0%, #764ba2 100%);
        color: white;
        border: none;
        border-radius: 25px;
        padding: 0.75rem 2rem;
        font-weight: 600;
        transition: all 0.3s ease;
    }
    
    .stButton > button:hover {
        transform: translateY(-2px);
        box-shadow: 0 5px 15px rgba(0,0,0,0.2);
    }
</style>
""", unsafe_allow_html=True)

def initialize_session_state():
    """Initialize session state variables"""
    if 'cv_data' not in st.session_state:
        st.session_state.cv_data = {
            'personal_info': {},
            'work_experience': [],
            'education': [],
            'skills': [],
            'languages': [],
            'certifications': [],
            'projects': []
        }
    
    if 'language' not in st.session_state:
        st.session_state.language = 'en'

def get_translations():
    """Get all translations for the application"""
    return {
        'en': {
            # Main titles
            'app_title': '📄 Professional CV Builder',
            'app_subtitle': '🚀 Create Your Professional CV in Minutes!',
            'app_description': 'Fill out the form below and generate a beautiful PDF or Word document instantly.',
            
            # Navigation
            'nav_personal': 'Personal Info',
            'nav_work': 'Work Experience', 
            'nav_education': 'Education',
            'nav_skills': 'Skills',
            'nav_languages': 'Languages',
            'nav_certifications': 'Certifications',
            'nav_projects': 'Projects',
            'nav_preview': 'Preview & Download',
            
            # Language selector
            'language_selector': 'Language / Dil',
            
            # Personal Info
            'personal_title': '👤 Personal Information',
            'first_name': 'First Name',
            'last_name': 'Last Name',
            'email': 'Email Address',
            'phone': 'Phone Number',
            'address': 'Address',
            'city': 'City',
            'country': 'Country',
            'date_birth': 'Date of Birth',
            'summary': 'Professional Summary',
            'summary_help': 'Write a brief summary of your professional background and career objectives.',
            
            # Work Experience
            'work_title': '💼 Work Experience',
            'company': 'Company Name',
            'position': 'Position',
            'start_date': 'Start Date',
            'end_date': 'End Date',
            'current_job': 'Current Job',
            'job_description': 'Job Description',
            
            # Education
            'education_title': '🎓 Education',
            'institution': 'Institution Name',
            'degree': 'Degree',
            'field_study': 'Field of Study',
            'gpa': 'GPA/Grade',
            
            # Skills
            'skills_title': '🛠️ Skills',
            'skill_name': 'Skill Name',
            'proficiency': 'Proficiency Level',
            'beginner': 'Beginner',
            'intermediate': 'Intermediate',
            'advanced': 'Advanced',
            'expert': 'Expert',
            
            # Languages
            'languages_title': '🌍 Languages',
            'language_name': 'Language Name',
            'fluency': 'Fluency Level',
            'basic': 'Basic',
            'conversational': 'Conversational',
            'fluent': 'Fluent',
            'native': 'Native',
            
            # Certifications
            'cert_title': '🏆 Certifications',
            'cert_name': 'Certificate Name',
            'issuing_org': 'Issuing Organization',
            'issue_date': 'Issue Date',
            'expiry_date': 'Expiry Date (Optional)',
            
            # Projects
            'projects_title': '🚀 Projects',
            'project_name': 'Project Name',
            'project_url': 'Project URL',
            'technologies': 'Technologies Used',
            'completion_date': 'Completion Date',
            'project_description': 'Project Description',
            
            # Buttons
            'add_entry': '➕ Add',
            'remove_entry': '➖ Remove Last Entry',
            'download_pdf': '📄 Download PDF',
            'download_word': '📝 Download Word',
            'click_download_pdf': '📄 Click to Download PDF',
            'click_download_word': '📝 Click to Download Word',
            
            # Preview & Download
            'preview_title': '👀 Preview & Download',
            'missing_info': '⚠️ Missing Required Information',
            'missing_desc': 'Please fill in the following required fields:',
            'complete_cv': '✅ CV Complete!',
            'complete_desc': 'Your CV looks great! Preview it below and download in your preferred format.',
            'cv_preview': '📋 CV Preview',
            'download_section': '📥 Download Your CV',
            
            # Messages
            'pdf_success': '✅ PDF generated successfully!',
            'word_success': '✅ Word document generated successfully!',
            'pdf_error': '❌ Error generating PDF:',
            'word_error': '❌ Error generating Word document:',
            'navigate_tip': '💡 Navigate to the respective sections using the sidebar to complete your CV.',
            
            # CV Sections in PDF/Word
            'cv_summary': 'Professional Summary',
            'cv_work': 'Work Experience',
            'cv_education': 'Education',
            'cv_skills': 'Skills',
            'cv_languages': 'Languages',
            'cv_certifications': 'Certifications',
            'cv_projects': 'Projects',
            
            # Footer
            'footer_title': '🚀 Professional CV Builder',
            'footer_desc': 'Create professional CVs in minutes with our modern, user-friendly interface!',
            
            # Required field indicators
            'required': '*',
            'at': 'at',
            'present': 'Present',
            'in': 'in'
        },
        'tr': {
            # Main titles
            'app_title': '📄 Profesyonel CV Oluşturucu',
            'app_subtitle': '🚀 Profesyonel CV\'nizi Dakikalar İçinde Oluşturun!',
            'app_description': 'Aşağıdaki formu doldurun ve anında güzel bir PDF veya Word belgesi oluşturun.',
            
            # Navigation
            'nav_personal': 'Kişisel Bilgiler',
            'nav_work': 'İş Deneyimi',
            'nav_education': 'Eğitim',
            'nav_skills': 'Yetenekler',
            'nav_languages': 'Diller',
            'nav_certifications': 'Sertifikalar',
            'nav_projects': 'Projeler',
            'nav_preview': 'Önizleme ve İndirme',
            
            # Language selector
            'language_selector': 'Dil / Language',
            
            # Personal Info
            'personal_title': '👤 Kişisel Bilgiler',
            'first_name': 'Ad',
            'last_name': 'Soyad',
            'email': 'E-posta Adresi',
            'phone': 'Telefon Numarası',
            'address': 'Adres',
            'city': 'Şehir',
            'country': 'Ülke',
            'date_birth': 'Doğum Tarihi',
            'summary': 'Profesyonel Özet',
            'summary_help': 'Profesyonel geçmişiniz ve kariyer hedefleriniz hakkında kısa bir özet yazın.',
            
            # Work Experience
            'work_title': '💼 İş Deneyimi',
            'company': 'Şirket Adı',
            'position': 'Pozisyon',
            'start_date': 'Başlangıç Tarihi',
            'end_date': 'Bitiş Tarihi',
            'current_job': 'Mevcut İş',
            'job_description': 'İş Açıklaması',
            
            # Education
            'education_title': '🎓 Eğitim',
            'institution': 'Kurum Adı',
            'degree': 'Derece',
            'field_study': 'Çalışma Alanı',
            'gpa': 'Not Ortalaması/Derece',
            
            # Skills
            'skills_title': '🛠️ Yetenekler',
            'skill_name': 'Yetenek Adı',
            'proficiency': 'Yeterlilik Seviyesi',
            'beginner': 'Başlangıç',
            'intermediate': 'Orta',
            'advanced': 'İleri',
            'expert': 'Uzman',
            
            # Languages
            'languages_title': '🌍 Diller',
            'language_name': 'Dil Adı',
            'fluency': 'Akıcılık Seviyesi',
            'basic': 'Temel',
            'conversational': 'Konuşma',
            'fluent': 'Akıcı',
            'native': 'Ana Dil',
            
            # Certifications
            'cert_title': '🏆 Sertifikalar',
            'cert_name': 'Sertifika Adı',
            'issuing_org': 'Veren Kuruluş',
            'issue_date': 'Veriliş Tarihi',
            'expiry_date': 'Son Geçerlilik Tarihi (İsteğe Bağlı)',
            
            # Projects
            'projects_title': '🚀 Projeler',
            'project_name': 'Proje Adı',
            'project_url': 'Proje URL\'si',
            'technologies': 'Kullanılan Teknolojiler',
            'completion_date': 'Tamamlanma Tarihi',
            'project_description': 'Proje Açıklaması',
            
            # Buttons
            'add_entry': '➕ Ekle',
            'remove_entry': '➖ Son Girişi Sil',
            'download_pdf': '📄 PDF İndir',
            'download_word': '📝 Word İndir',
            'click_download_pdf': '📄 PDF İndirmek İçin Tıklayın',
            'click_download_word': '📝 Word İndirmek İçin Tıklayın',
            
            # Preview & Download
            'preview_title': '👀 Önizleme ve İndirme',
            'missing_info': '⚠️ Eksik Zorunlu Bilgiler',
            'missing_desc': 'Lütfen aşağıdaki zorunlu alanları doldurun:',
            'complete_cv': '✅ CV Tamamlandı!',
            'complete_desc': 'CV\'niz harika görünüyor! Aşağıda önizleyin ve tercih ettiğiniz formatta indirin.',
            'cv_preview': '📋 CV Önizlemesi',
            'download_section': '📥 CV\'nizi İndirin',
            
            # Messages
            'pdf_success': '✅ PDF başarıyla oluşturuldu!',
            'word_success': '✅ Word belgesi başarıyla oluşturuldu!',
            'pdf_error': '❌ PDF oluşturma hatası:',
            'word_error': '❌ Word belgesi oluşturma hatası:',
            'navigate_tip': '💡 CV\'nizi tamamlamak için yan menüyü kullanarak ilgili bölümlere gidin.',
            
            # CV Sections in PDF/Word
            'cv_summary': 'Profesyonel Özet',
            'cv_work': 'İş Deneyimi',
            'cv_education': 'Eğitim',
            'cv_skills': 'Yetenekler',
            'cv_languages': 'Diller',
            'cv_certifications': 'Sertifikalar',
            'cv_projects': 'Projeler',
            
            # Footer
            'footer_title': '🚀 Profesyonel CV Oluşturucu',
            'footer_desc': 'Modern ve kullanıcı dostu arayüzümüzle dakikalar içinde profesyonel CV\'ler oluşturun!',
            
            # Required field indicators
            'required': '*',
            'at': 'şirketinde',
            'present': 'Devam Ediyor',
            'in': 'alanında'
        }
    }

def t(key, lang=None):
    """Get translation for given key"""
    if lang is None:
        lang = st.session_state.get('language', 'en')
    
    translations = get_translations()
    return translations.get(lang, {}).get(key, key)

def validate_required_fields(data):
    """Validate required fields and return missing fields"""
    missing_fields = []
    
    # Check personal info
    required_personal = ['first_name', 'last_name', 'email', 'phone']
    for field in required_personal:
        if not data['personal_info'].get(field, '').strip():
            missing_fields.append(f"Personal Info: {field.replace('_', ' ').title()}")
    
    # Check if at least one work experience exists
    if not data['work_experience'] or not any(exp.get('company', '').strip() for exp in data['work_experience']):
        missing_fields.append("At least one work experience")
    
    # Check if at least one education exists
    if not data['education'] or not any(edu.get('institution', '').strip() for edu in data['education']):
        missing_fields.append("At least one education entry")
    
    return missing_fields

def create_pdf(data, lang='en'):
    """Create PDF document from CV data"""
    buffer = io.BytesIO()
    doc = SimpleDocTemplate(buffer, pagesize=A4, topMargin=0.5*inch)
    styles = getSampleStyleSheet()
    story = []
    
    # Custom styles
    title_style = ParagraphStyle(
        'CustomTitle',
        parent=styles['Heading1'],
        fontSize=24,
        spaceAfter=20,
        alignment=1,
        textColor=colors.HexColor('#2c3e50')
    )
    
    heading_style = ParagraphStyle(
        'CustomHeading',
        parent=styles['Heading2'],
        fontSize=16,
        spaceAfter=12,
        textColor=colors.HexColor('#3498db'),
        borderWidth=1,
        borderColor=colors.HexColor('#3498db'),
        borderPadding=5
    )
    
    # Title
    full_name = f"{data['personal_info'].get('first_name', '')} {data['personal_info'].get('last_name', '')}"
    story.append(Paragraph(full_name, title_style))
    
    # Contact info
    contact_info = f"""
    <b>{t('email', lang)}:</b> {data['personal_info'].get('email', '')}<br/>
    <b>{t('phone', lang)}:</b> {data['personal_info'].get('phone', '')}<br/>
    <b>{t('address', lang)}:</b> {data['personal_info'].get('address', '')}<br/>
    <b>{t('city', lang)}:</b> {data['personal_info'].get('city', '')}, {data['personal_info'].get('country', '')}
    """
    story.append(Paragraph(contact_info, styles['Normal']))
    story.append(Spacer(1, 20))
    
    # Profile Summary
    if data['personal_info'].get('summary'):
        story.append(Paragraph(t('cv_summary', lang), heading_style))
        story.append(Paragraph(data['personal_info']['summary'], styles['Normal']))
        story.append(Spacer(1, 15))
    
    # Work Experience
    if data['work_experience']:
        story.append(Paragraph(t('cv_work', lang), heading_style))
        for work in data['work_experience']:
            if work.get('company'):
                work_title = f"<b>{work.get('position', '')}</b> {t('at', lang)} <b>{work.get('company', '')}</b>"
                story.append(Paragraph(work_title, styles['Normal']))
                
                end_date_text = work.get('end_date', t('present', lang) if work.get('current_job') else '')
                dates = f"{work.get('start_date', '')} - {end_date_text}"
                story.append(Paragraph(dates, styles['Normal']))
                
                if work.get('description'):
                    story.append(Paragraph(work['description'], styles['Normal']))
                story.append(Spacer(1, 10))
    
    # Education
    if data['education']:
        story.append(Paragraph(t('cv_education', lang), heading_style))
        for edu in data['education']:
            if edu.get('institution'):
                edu_title = f"<b>{edu.get('degree', '')}</b> {t('in', lang)} {edu.get('field', '')}"
                story.append(Paragraph(edu_title, styles['Normal']))
                story.append(Paragraph(edu['institution'], styles['Normal']))
                
                dates = f"{edu.get('start_date', '')} - {edu.get('end_date', '')}"
                story.append(Paragraph(dates, styles['Normal']))
                
                if edu.get('gpa'):
                    gpa_label = t('gpa', lang) if lang == 'en' else 'Not Ortalaması'
                    story.append(Paragraph(f"{gpa_label}: {edu['gpa']}", styles['Normal']))
                story.append(Spacer(1, 10))
    
    # Skills
    if data['skills']:
        story.append(Paragraph(t('cv_skills', lang), heading_style))
        skills_text = ""
        for skill in data['skills']:
            if skill.get('name'):
                skills_text += f"• {skill['name']} ({skill.get('level', t('intermediate', lang))})<br/>"
        if skills_text:
            story.append(Paragraph(skills_text, styles['Normal']))
            story.append(Spacer(1, 15))
    
    # Languages
    if data['languages']:
        story.append(Paragraph(t('cv_languages', lang), heading_style))
        lang_text = ""
        for language in data['languages']:
            if language.get('name'):
                lang_text += f"• {language['name']} ({language.get('level', t('conversational', lang))})<br/>"
        if lang_text:
            story.append(Paragraph(lang_text, styles['Normal']))
            story.append(Spacer(1, 15))
    
    # Certifications
    if data['certifications']:
        story.append(Paragraph(t('cv_certifications', lang), heading_style))
        for cert in data['certifications']:
            if cert.get('name'):
                cert_title = f"<b>{cert['name']}</b>"
                story.append(Paragraph(cert_title, styles['Normal']))
                if cert.get('issuer'):
                    story.append(Paragraph(cert['issuer'], styles['Normal']))
                if cert.get('issue_date'):
                    story.append(Paragraph(cert['issue_date'], styles['Normal']))
                story.append(Spacer(1, 10))
    
    # Projects
    if data['projects']:
        story.append(Paragraph(t('cv_projects', lang), heading_style))
        for project in data['projects']:
            if project.get('name'):
                proj_title = f"<b>{project['name']}</b>"
                story.append(Paragraph(proj_title, styles['Normal']))
                if project.get('technologies'):
                    tech_label = t('technologies', lang)
                    story.append(Paragraph(f"{tech_label}: {project['technologies']}", styles['Normal']))
                if project.get('description'):
                    story.append(Paragraph(project['description'], styles['Normal']))
                story.append(Spacer(1, 10))
    
    doc.build(story)
    buffer.seek(0)
    return buffer

def create_word_doc(data, lang='en'):
    """Create Word document from CV data"""
    doc = Document()
    
    # Title
    title = doc.add_heading('', 0)
    full_name = f"{data['personal_info'].get('first_name', '')} {data['personal_info'].get('last_name', '')}"
    title_run = title.runs[0] if title.runs else title.add_run()
    title_run.text = full_name
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Contact info
    contact_para = doc.add_paragraph()
    contact_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    contact_text = f"""{t('email', lang)}: {data['personal_info'].get('email', '')}
{t('phone', lang)}: {data['personal_info'].get('phone', '')}
{t('address', lang)}: {data['personal_info'].get('address', '')}
{data['personal_info'].get('city', '')}, {data['personal_info'].get('country', '')}"""
    contact_para.add_run(contact_text)
    
    # Profile Summary
    if data['personal_info'].get('summary'):
        doc.add_heading(t('cv_summary', lang), level=1)
        doc.add_paragraph(data['personal_info']['summary'])
    
    # Work Experience
    if data['work_experience']:
        doc.add_heading(t('cv_work', lang), level=1)
        for work in data['work_experience']:
            if work.get('company'):
                work_para = doc.add_paragraph()
                work_para.add_run(f"{work.get('position', '')} {t('at', lang)} {work.get('company', '')}").bold = True
                
                end_date_text = work.get('end_date', t('present', lang) if work.get('current_job') else '')
                date_para = doc.add_paragraph(f"{work.get('start_date', '')} - {end_date_text}")
                
                if work.get('description'):
                    doc.add_paragraph(work['description'])
    
    # Education
    if data['education']:
        doc.add_heading(t('cv_education', lang), level=1)
        for edu in data['education']:
            if edu.get('institution'):
                edu_para = doc.add_paragraph()
                edu_para.add_run(f"{edu.get('degree', '')} {t('in', lang)} {edu.get('field', '')}").bold = True
                
                doc.add_paragraph(edu['institution'])
                doc.add_paragraph(f"{edu.get('start_date', '')} - {edu.get('end_date', '')}")
                
                if edu.get('gpa'):
                    gpa_label = t('gpa', lang) if lang == 'en' else 'Not Ortalaması'
                    doc.add_paragraph(f"{gpa_label}: {edu['gpa']}")
    
    # Skills
    if data['skills']:
        doc.add_heading(t('cv_skills', lang), level=1)
        for skill in data['skills']:
            if skill.get('name'):
                doc.add_paragraph(f"• {skill['name']} ({skill.get('level', t('intermediate', lang))})")
    
    # Languages
    if data['languages']:
        doc.add_heading(t('cv_languages', lang), level=1)
        for language in data['languages']:
            if language.get('name'):
                doc.add_paragraph(f"• {language['name']} ({language.get('level', t('conversational', lang))})")
    
    # Certifications
    if data['certifications']:
        doc.add_heading(t('cv_certifications', lang), level=1)
        for cert in data['certifications']:
            if cert.get('name'):
                cert_para = doc.add_paragraph()
                cert_para.add_run(cert['name']).bold = True
                if cert.get('issuer'):
                    doc.add_paragraph(cert['issuer'])
                if cert.get('issue_date'):
                    doc.add_paragraph(cert['issue_date'])
    
    # Projects
    if data['projects']:
        doc.add_heading(t('cv_projects', lang), level=1)
        for project in data['projects']:
            if project.get('name'):
                proj_para = doc.add_paragraph()
                proj_para.add_run(project['name']).bold = True
                if project.get('technologies'):
                    tech_label = t('technologies', lang)
                    doc.add_paragraph(f"{tech_label}: {project['technologies']}")
                if project.get('description'):
                    doc.add_paragraph(project['description'])
    
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer

def display_cv_preview(data, lang='en'):
    """Display CV preview"""
    st.markdown('<div class="preview-container">', unsafe_allow_html=True)
    
    # Header
    full_name = f"{data['personal_info'].get('first_name', '')} {data['personal_info'].get('last_name', '')}"
    st.markdown(f"<h1 style='text-align: center; color: #2c3e50;'>{full_name}</h1>", unsafe_allow_html=True)
    
    # Contact Info
    contact_info = f"""
    **{t('email', lang)}:** {data['personal_info'].get('email', '')} | 
    **{t('phone', lang)}:** {data['personal_info'].get('phone', '')} | 
    **{t('address', lang)}:** {data['personal_info'].get('address', '')}, {data['personal_info'].get('city', '')}, {data['personal_info'].get('country', '')}
    """
    st.markdown(f"<p style='text-align: center;'>{contact_info}</p>", unsafe_allow_html=True)
    
    # Profile Summary
    if data['personal_info'].get('summary'):
        st.markdown(f"### 🎯 {t('cv_summary', lang)}")
        st.write(data['personal_info']['summary'])
    
    # Work Experience
    if data['work_experience']:
        st.markdown(f"### 💼 {t('cv_work', lang)}")
        for work in data['work_experience']:
            if work.get('company'):
                st.markdown(f"**{work.get('position', '')}** {t('at', lang)} **{work.get('company', '')}**")
                end_date_text = work.get('end_date', t('present', lang) if work.get('current_job') else '')
                st.markdown(f"*{work.get('start_date', '')} - {end_date_text}*")
                if work.get('description'):
                    st.write(work['description'])
                st.markdown("---")
    
    # Education
    if data['education']:
        st.markdown(f"### 🎓 {t('cv_education', lang)}")
        for edu in data['education']:
            if edu.get('institution'):
                st.markdown(f"**{edu.get('degree', '')}** {t('in', lang)} {edu.get('field', '')}")
                st.markdown(f"*{edu['institution']}*")
                st.markdown(f"*{edu.get('start_date', '')} - {edu.get('end_date', '')}*")
                if edu.get('gpa'):
                    gpa_label = t('gpa', lang) if lang == 'en' else 'Not Ortalaması'
                    st.markdown(f"**{gpa_label}:** {edu['gpa']}")
                st.markdown("---")
    
    # Skills
    if data['skills']:
        st.markdown(f"### 🛠️ {t('cv_skills', lang)}")
        for skill in data['skills']:
            if skill.get('name'):
                st.markdown(f"• **{skill['name']}** ({skill.get('level', t('intermediate', lang))})")
    
    # Languages
    if data['languages']:
        st.markdown(f"### 🌍 {t('cv_languages', lang)}")
        for language in data['languages']:
            if language.get('name'):
                st.markdown(f"• **{language['name']}** ({language.get('level', t('conversational', lang))})")
    
    # Certifications
    if data['certifications']:
        st.markdown(f"### 🏆 {t('cv_certifications', lang)}")
        for cert in data['certifications']:
            if cert.get('name'):
                st.markdown(f"• **{cert['name']}**")
                if cert.get('issuer'):
                    st.markdown(f"  *{cert['issuer']}*")
                if cert.get('issue_date'):
                    st.markdown(f"  *{cert['issue_date']}*")
    
    # Projects
    if data['projects']:
        st.markdown(f"### 🚀 {t('cv_projects', lang)}")
        for project in data['projects']:
            if project.get('name'):
                st.markdown(f"• **{project['name']}**")
                if project.get('technologies'):
                    tech_label = t('technologies', lang)
                    st.markdown(f"  *{tech_label}: {project['technologies']}*")
                if project.get('description'):
                    st.markdown(f"  {project['description']}")
    
    st.markdown('</div>', unsafe_allow_html=True)

def main():
    """Main application function"""
    initialize_session_state()
    
    # Language selector at the top
    col1, col2, col3 = st.columns([2, 1, 2])
    with col2:
        language_options = {'English': 'en', 'Türkçe': 'tr'}
        selected_lang = st.selectbox(
            t('language_selector'),
            options=list(language_options.keys()),
            index=0 if st.session_state.language == 'en' else 1,
            key='lang_selector'
        )
        
        # Update language in session state
        new_lang = language_options[selected_lang]
        if new_lang != st.session_state.language:
            st.session_state.language = new_lang
            st.rerun()
    
    # Header
    st.markdown(f'<h1 class="main-header">{t("app_title")}</h1>', unsafe_allow_html=True)
    st.markdown(f'<div class="info-box"><h3>{t("app_subtitle")}</h3><p>{t("app_description")}</p></div>', unsafe_allow_html=True)
    
    # Sidebar for navigation
    st.sidebar.title("📋 CV Sections")
    sections = [
        t('nav_personal'), t('nav_work'), t('nav_education'), 
        t('nav_skills'), t('nav_languages'), t('nav_certifications'), 
        t('nav_projects'), t('nav_preview')
    ]
    selected_section = st.sidebar.radio("Navigate to:", sections)
    
    # Main content area
    if selected_section == t('nav_personal'):
        st.markdown(f'<h2 class="section-header">{t("personal_title")}</h2>', unsafe_allow_html=True)
        
        col1, col2 = st.columns(2)
        
        with col1:
            first_name = st.text_input(f"{t('first_name')} {t('required')}", value=st.session_state.cv_data['personal_info'].get('first_name', ''))
            email = st.text_input(f"{t('email')} {t('required')}", value=st.session_state.cv_data['personal_info'].get('email', ''))
            address = st.text_area(t('address'), value=st.session_state.cv_data['personal_info'].get('address', ''))
            country = st.text_input(t('country'), value=st.session_state.cv_data['personal_info'].get('country', ''))
        
        with col2:
            last_name = st.text_input(f"{t('last_name')} {t('required')}", value=st.session_state.cv_data['personal_info'].get('last_name', ''))
            phone = st.text_input(f"{t('phone')} {t('required')}", value=st.session_state.cv_data['personal_info'].get('phone', ''))
            city = st.text_input(t('city'), value=st.session_state.cv_data['personal_info'].get('city', ''))
            date_of_birth = st.date_input(t('date_birth'), value=date(1990, 1, 1))
        
        summary = st.text_area(t('summary'), 
                              value=st.session_state.cv_data['personal_info'].get('summary', ''),
                              help=t('summary_help'))
        
        # Save to session state
        st.session_state.cv_data['personal_info'] = {
            'first_name': first_name,
            'last_name': last_name,
            'email': email,
            'phone': phone,
            'address': address,
            'city': city,
            'country': country,
            'date_of_birth': str(date_of_birth),
            'summary': summary
        }
    
    elif selected_section == t('nav_work'):
        st.markdown(f'<h2 class="section-header">{t("work_title")}</h2>', unsafe_allow_html=True)
        
        # Initialize work experience list if empty
        if not st.session_state.cv_data['work_experience']:
            st.session_state.cv_data['work_experience'] = [{}]
        
        for i, work in enumerate(st.session_state.cv_data['work_experience']):
            with st.expander(f"{t('nav_work')} #{i+1}", expanded=True):
                col1, col2 = st.columns(2)
                
                with col1:
                    company = st.text_input(f"{t('company')} {t('required')}", key=f"work_company_{i}", value=work.get('company', ''))
                    start_date = st.date_input(t('start_date'), key=f"work_start_{i}", value=date(2020, 1, 1))
                    current_job = st.checkbox(t('current_job'), key=f"work_current_{i}", value=work.get('current_job', False))
                
                with col2:
                    position = st.text_input(f"{t('position')} {t('required')}", key=f"work_position_{i}", value=work.get('position', ''))
                    if not current_job:
                        end_date = st.date_input(t('end_date'), key=f"work_end_{i}", value=date(2023, 12, 31))
                    else:
                        end_date = None
                
                description = st.text_area(t('job_description'), key=f"work_desc_{i}", value=work.get('description', ''))
                
                # Update session state
                st.session_state.cv_data['work_experience'][i] = {
                    'company': company,
                    'position': position,
                    'start_date': str(start_date),
                    'end_date': str(end_date) if end_date else '',
                    'current_job': current_job,
                    'description': description
                }
        
        col1, col2 = st.columns(2)
        with col1:
            if st.button(f"{t('add_entry')} {t('nav_work')}"):
                st.session_state.cv_data['work_experience'].append({})
                st.rerun()
        
        with col2:
            if len(st.session_state.cv_data['work_experience']) > 1:
                if st.button(t('remove_entry')):
                    st.session_state.cv_data['work_experience'].pop()
                    st.rerun()
    
    elif selected_section == t('nav_education'):
        st.markdown(f'<h2 class="section-header">{t("education_title")}</h2>', unsafe_allow_html=True)
        
        # Initialize education list if empty
        if not st.session_state.cv_data['education']:
            st.session_state.cv_data['education'] = [{}]
        
        for i, edu in enumerate(st.session_state.cv_data['education']):
            with st.expander(f"{t('nav_education')} #{i+1}", expanded=True):
                col1, col2 = st.columns(2)
                
                with col1:
                    institution = st.text_input(f"{t('institution')} {t('required')}", key=f"edu_institution_{i}", value=edu.get('institution', ''))
                    field = st.text_input(t('field_study'), key=f"edu_field_{i}", value=edu.get('field', ''))
                    start_date = st.date_input(t('start_date'), key=f"edu_start_{i}", value=date(2018, 9, 1))
                
                with col2:
                    degree = st.text_input(f"{t('degree')} {t('required')}", key=f"edu_degree_{i}", value=edu.get('degree', ''))
                    gpa = st.text_input(t('gpa'), key=f"edu_gpa_{i}", value=edu.get('gpa', ''))
                    end_date = st.date_input(t('end_date'), key=f"edu_end_{i}", value=date(2022, 6, 30))
                
                # Update session state
                st.session_state.cv_data['education'][i] = {
                    'institution': institution,
                    'degree': degree,
                    'field': field,
                    'start_date': str(start_date),
                    'end_date': str(end_date),
                    'gpa': gpa
                }
        
        col1, col2 = st.columns(2)
        with col1:
            if st.button(f"{t('add_entry')} {t('nav_education')}"):
                st.session_state.cv_data['education'].append({})
                st.rerun()
        
        with col2:
            if len(st.session_state.cv_data['education']) > 1:
                if st.button(t('remove_entry')):
                    st.session_state.cv_data['education'].pop()
                    st.rerun()
    
    elif selected_section == t('nav_skills'):
        st.markdown(f'<h2 class="section-header">{t("skills_title")}</h2>', unsafe_allow_html=True)
        
        # Initialize skills list if empty
        if not st.session_state.cv_data['skills']:
            st.session_state.cv_data['skills'] = [{}]
        
        for i, skill in enumerate(st.session_state.cv_data['skills']):
            with st.expander(f"{t('nav_skills')} #{i+1}", expanded=True):
                col1, col2 = st.columns(2)
                
                with col1:
                    name = st.text_input(f"{t('skill_name')} {t('required')}", key=f"skill_name_{i}", value=skill.get('name', ''))
                
                with col2:
                    level_options = [t('beginner'), t('intermediate'), t('advanced'), t('expert')]
                    current_level = skill.get('level', t('intermediate'))
                    try:
                        level_index = level_options.index(current_level)
                    except ValueError:
                        level_index = 1  # Default to intermediate
                    
                    level = st.selectbox(t('proficiency'), 
                                       level_options,
                                       key=f"skill_level_{i}",
                                       index=level_index)
                
                # Update session state
                st.session_state.cv_data['skills'][i] = {
                    'name': name,
                    'level': level
                }
        
        col1, col2 = st.columns(2)
        with col1:
            if st.button(f"{t('add_entry')} {t('nav_skills')}"):
                st.session_state.cv_data['skills'].append({})
                st.rerun()
        
        with col2:
            if len(st.session_state.cv_data['skills']) > 1:
                if st.button(t('remove_entry')):
                    st.session_state.cv_data['skills'].pop()
                    st.rerun()
    
    elif selected_section == t('nav_languages'):
        st.markdown(f'<h2 class="section-header">{t("languages_title")}</h2>', unsafe_allow_html=True)
        
        # Initialize languages list if empty
        if not st.session_state.cv_data['languages']:
            st.session_state.cv_data['languages'] = [{}]
        
        for i, lang in enumerate(st.session_state.cv_data['languages']):
            with st.expander(f"{t('nav_languages')} #{i+1}", expanded=True):
                col1, col2 = st.columns(2)
                
                with col1:
                    name = st.text_input(f"{t('language_name')} {t('required')}", key=f"lang_name_{i}", value=lang.get('name', ''))
                
                with col2:
                    level_options = [t('basic'), t('conversational'), t('fluent'), t('native')]
                    current_level = lang.get('level', t('conversational'))
                    try:
                        level_index = level_options.index(current_level)
                    except ValueError:
                        level_index = 1  # Default to conversational
                    
                    level = st.selectbox(t('fluency'), 
                                       level_options,
                                       key=f"lang_level_{i}",
                                       index=level_index)
                
                # Update session state
                st.session_state.cv_data['languages'][i] = {
                    'name': name,
                    'level': level
                }
        
        col1, col2 = st.columns(2)
        with col1:
            if st.button(f"{t('add_entry')} {t('nav_languages')}"):
                st.session_state.cv_data['languages'].append({})
                st.rerun()
        
        with col2:
            if len(st.session_state.cv_data['languages']) > 1:
                if st.button(t('remove_entry')):
                    st.session_state.cv_data['languages'].pop()
                    st.rerun()
    
    elif selected_section == t('nav_certifications'):
        st.markdown(f'<h2 class="section-header">{t("cert_title")}</h2>', unsafe_allow_html=True)
        
        # Initialize certifications list if empty
        if not st.session_state.cv_data['certifications']:
            st.session_state.cv_data['certifications'] = [{}]
        
        for i, cert in enumerate(st.session_state.cv_data['certifications']):
            with st.expander(f"{t('nav_certifications')} #{i+1}", expanded=True):
                col1, col2 = st.columns(2)
                
                with col1:
                    name = st.text_input(t('cert_name'), key=f"cert_name_{i}", value=cert.get('name', ''))
                    issuer = st.text_input(t('issuing_org'), key=f"cert_issuer_{i}", value=cert.get('issuer', ''))
                
                with col2:
                    issue_date = st.date_input(t('issue_date'), key=f"cert_date_{i}", value=date(2023, 1, 1))
                    expiry_date = st.date_input(t('expiry_date'), key=f"cert_expiry_{i}", value=None)
                
                # Update session state
                st.session_state.cv_data['certifications'][i] = {
                    'name': name,
                    'issuer': issuer,
                    'issue_date': str(issue_date),
                    'expiry_date': str(expiry_date) if expiry_date else ''
                }
        
        col1, col2 = st.columns(2)
        with col1:
            if st.button(f"{t('add_entry')} {t('nav_certifications')}"):
                st.session_state.cv_data['certifications'].append({})
                st.rerun()
        
        with col2:
            if len(st.session_state.cv_data['certifications']) > 1:
                if st.button(t('remove_entry')):
                    st.session_state.cv_data['certifications'].pop()
                    st.rerun()
    
    elif selected_section == t('nav_projects'):
        st.markdown(f'<h2 class="section-header">{t("projects_title")}</h2>', unsafe_allow_html=True)
        
        # Initialize projects list if empty
        if not st.session_state.cv_data['projects']:
            st.session_state.cv_data['projects'] = [{}]
        
        for i, project in enumerate(st.session_state.cv_data['projects']):
            with st.expander(f"{t('nav_projects')} #{i+1}", expanded=True):
                col1, col2 = st.columns(2)
                
                with col1:
                    name = st.text_input(t('project_name'), key=f"proj_name_{i}", value=project.get('name', ''))
                    technologies = st.text_input(t('technologies'), key=f"proj_tech_{i}", value=project.get('technologies', ''))
                
                with col2:
                    url = st.text_input(t('project_url'), key=f"proj_url_{i}", value=project.get('url', ''))
                    date_completed = st.date_input(t('completion_date'), key=f"proj_date_{i}", value=date(2023, 6, 1))
                
                description = st.text_area(t('project_description'), key=f"proj_desc_{i}", value=project.get('description', ''))
                
                # Update session state
                st.session_state.cv_data['projects'][i] = {
                    'name': name,
                    'url': url,
                    'technologies': technologies,
                    'date_completed': str(date_completed),
                    'description': description
                }
        
        col1, col2 = st.columns(2)
        with col1:
            if st.button(f"{t('add_entry')} {t('nav_projects')}"):
                st.session_state.cv_data['projects'].append({})
                st.rerun()
        
        with col2:
            if len(st.session_state.cv_data['projects']) > 1:
                if st.button(t('remove_entry')):
                    st.session_state.cv_data['projects'].pop()
                    st.rerun()
    
    elif selected_section == t('nav_preview'):
        st.markdown(f'<h2 class="section-header">{t("preview_title")}</h2>', unsafe_allow_html=True)
        
        # Validate required fields
        missing_fields = validate_required_fields(st.session_state.cv_data)
        
        if missing_fields:
            st.markdown(f'<div class="warning-box"><h3>{t("missing_info")}</h3><p>{t("missing_desc")}</p></div>', unsafe_allow_html=True)
            for field in missing_fields:
                st.error(f"❌ {field}")
            st.info(t('navigate_tip'))
        else:
            st.markdown(f'<div class="success-box"><h3>{t("complete_cv")}</h3><p>{t("complete_desc")}</p></div>', unsafe_allow_html=True)
            
            # Preview
            st.markdown(f"## {t('cv_preview')}")
            display_cv_preview(st.session_state.cv_data, st.session_state.language)
            
            # Download buttons
            st.markdown(f"## {t('download_section')}")
            col1, col2, col3 = st.columns([1, 1, 2])
            
            with col1:
                if st.button(t('download_pdf'), type="primary"):
                    try:
                        pdf_buffer = create_pdf(st.session_state.cv_data, st.session_state.language)
                        st.download_button(
                            label=t('click_download_pdf'),
                            data=pdf_buffer,
                            file_name=f"CV_{st.session_state.cv_data['personal_info']['first_name']}_{st.session_state.cv_data['personal_info']['last_name']}.pdf",
                            mime="application/pdf"
                        )
                        st.success(t('pdf_success'))
                    except Exception as e:
                        st.error(f"{t('pdf_error')} {str(e)}")
            
            with col2:
                if st.button(t('download_word'), type="primary"):
                    try:
                        word_buffer = create_word_doc(st.session_state.cv_data, st.session_state.language)
                        st.download_button(
                            label=t('click_download_word'),
                            data=word_buffer,
                            file_name=f"CV_{st.session_state.cv_data['personal_info']['first_name']}_{st.session_state.cv_data['personal_info']['last_name']}.docx",
                            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                        )
                        st.success(t('word_success'))
                    except Exception as e:
                        st.error(f"{t('word_error')} {str(e)}")
    
    # Footer
    st.markdown("---")
    st.markdown(
        f"""
        <div style='text-align: center; color: #7f8c8d; padding: 2rem;'>
            <p>🚀 <strong>{t('footer_title')}</strong> | Built with ❤️ using Streamlit</p>
            <p>{t('footer_desc')}</p>
        </div>
        """, 
        unsafe_allow_html=True
    )

if __name__ == "__main__":
    main()
